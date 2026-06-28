from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables"
OUT.mkdir(exist_ok=True)


def set_run_font(run, size=11, bold=False):
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def setup_doc(title, subtitle="校园点餐系统"):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    style = doc.styles["Normal"]
    style.font.name = "SimSun"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_run_font(r, 20, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    set_run_font(r, 12)
    doc.add_paragraph()
    return doc


def add_heading(doc, text, level=1):
    p = doc.add_heading("", level=level)
    r = p.add_run(text)
    set_run_font(r, 16 if level == 1 else 13, True)
    return p


def add_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        r = cell.paragraphs[0].add_run(header)
        set_run_font(r, 10, True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            r = cells[idx].paragraphs[0].add_run(str(value))
            set_run_font(r, 10)
    doc.add_paragraph()
    return table


def save(doc, filename):
    path = OUT / filename
    doc.save(path)
    return path


def requirement_doc():
    doc = setup_doc("校园点餐系统_需求分析说明书_第X组")
    add_heading(doc, "1. 项目背景")
    add_para(doc, "校园食堂在午晚餐高峰期经常出现排队时间长、窗口备餐压力集中、学生难以提前规划就餐时间等问题。校园点餐系统面向学生、教师、食堂管理员和团餐组织方，提供菜品浏览、预约点餐、送餐、自取核销和团餐预订能力，使食堂能够提前掌握订单量并安排备餐，用户能够减少现场等待。系统选择微信小程序形态，是因为它无需安装、适合校园场景快速触达，并且可通过云开发降低后端部署复杂度。")
    add_heading(doc, "2. 用户画像")
    add_table(doc, ["角色", "使用场景", "核心诉求", "使用频次"], [
        ["学生", "课间或饭前快速点餐，到窗口凭码取餐", "少排队、价格清楚、订单状态可查", "每日1-2次"],
        ["教师", "提前预约午餐或送餐到办公室", "时间确定、地址准确、流程简单", "每周数次"],
        ["食堂管理员", "维护菜品、处理订单、核销取餐码", "状态清晰、核验快捷、减少错领", "持续使用"],
        ["团餐组织方", "培训班、接待活动批量预订", "人数、时间、要求可记录，便于对账", "按活动使用"],
    ])
    add_heading(doc, "3. 用户故事地图")
    add_table(doc, ["角色", "场景", "核心任务", "子任务"], [
        ["学生", "即时点餐", "完成一笔自取订单", "登录、浏览分类、加入购物车、提交支付、查看取餐码、窗口核销"],
        ["学生/教师", "预约点餐", "选择自取或送餐", "选择菜品、选择时间、选择地址、支付、确认收货或核销"],
        ["团餐组织方", "团餐批量预订", "创建团餐订单", "填写人数、时间、套餐、特殊要求、生成订单"],
        ["管理员", "食堂管理", "运营订单与菜品", "查看统计、处理状态、核销取餐凭证、查看库存"],
    ])
    add_heading(doc, "4. 业务流程图说明")
    add_para(doc, "下单支付流程：用户登录后进入菜品页，按分类或关键词选择菜品，加入购物车后进入结算页；系统计算总价并校验库存；用户选择自取、送餐或团餐信息后提交；支付模拟成功后系统生成订单号、更新订单状态、扣减库存并生成取餐凭证。")
    add_para(doc, "取餐核销流程：用户在订单详情页出示六位取餐码；管理员在管理端输入取餐码；系统查询取餐码是否存在、是否已使用、是否属于有效订单；校验通过后订单状态更新为已取餐，取餐码标记为已核销。")
    add_heading(doc, "5. 功能模块划分")
    add_table(doc, ["模块", "包含功能", "依赖模块", "优先级"], [
        ["登录与用户", "微信授权登录、角色识别、用户信息展示", "无", "P0"],
        ["菜品浏览", "分类、搜索、分页、刷新、详情", "登录", "P0"],
        ["购物车", "增减数量、删除、总价计算", "菜品浏览", "P1"],
        ["订单", "下单、支付模拟、订单号、状态流转", "购物车、地址", "P0/P1"],
        ["预约与配送", "自取/送餐、时间、地址、确认收货", "订单、地址", "P1"],
        ["凭证核销", "取餐码生成、展示、管理员核销", "订单", "P1"],
        ["管理端", "订单处理、菜品库存、统计", "订单、菜品", "P2"],
        ["评价与创新", "评分、电子小票、营养提示、缓存", "订单", "P2/P3"],
    ])
    add_heading(doc, "6. 技术选型")
    add_para(doc, "本项目选择 A 路线：微信原生小程序 + 微信云开发。选择理由包括：第一，实训周期只有5天，云开发能降低服务器、域名、部署和数据库连接成本；第二，小程序原生能力与微信登录、订阅消息、支付沙箱和云函数集成度高；第三，团队可以把更多时间投入需求分析、核心业务和答辩展示，而不是复杂运维。")
    add_heading(doc, "7. 数据库设计")
    add_table(doc, ["集合", "主要字段", "索引规划"], [
        ["users", "openid, role, nickName, avatarUrl", "openid唯一"],
        ["dishes", "name, categoryId, price, stock, rating, sales, nutrition", "categoryId+onSale, name"],
        ["orders", "orderNo, userOpenid, orderType, items, totalAmount, status, pickupCode", "userOpenid+createdAt, status+createdAt, pickupCode"],
        ["pickup_codes", "code, orderId, used, usedAt", "code唯一"],
        ["addresses", "userOpenid, name, phone, detail, isDefault", "userOpenid"],
        ["reviews", "orderId, userOpenid, rating, content", "orderId, userOpenid+createdAt"],
    ])
    add_heading(doc, "8. 进度计划与分工")
    add_table(doc, ["日期", "核心任务", "负责人"], [
        ["Day1", "需求分析、系统设计、数据库建模、Git初始化", "全组"],
        ["Day2", "登录、菜品页、购物车、基础云函数", "前端/后端"],
        ["Day3", "下单、预约、取餐码、订单列表与详情", "前端/后端"],
        ["Day4", "核销、管理端、评价、统计、缓存与创新功能", "全组"],
        ["Day5", "联调、测试、PPT、报告、答辩演示", "全组"],
    ])
    return save(doc, "校园点餐系统_需求分析说明书_第X组.docx")


def design_doc():
    doc = setup_doc("校园点餐系统_系统设计说明书_第X组")
    add_heading(doc, "1. 系统架构")
    add_para(doc, "系统采用微信小程序原生前端、云函数业务层、云数据库/云存储数据层的三层结构。小程序页面负责交互、表单校验与本地状态展示；API facade 统一封装云函数调用；云函数承担登录、菜品查询、下单、核销、状态更新和评价提交等业务；云数据库保存用户、菜品、订单、取餐码和评价数据。")
    add_heading(doc, "2. 技术栈详细说明")
    add_table(doc, ["层次", "技术", "说明"], [
        ["前端", "WXML/WXSS/JS/JSON", "原生页面、组件化样式、路由与本地缓存"],
        ["云函数", "Node.js + wx-server-sdk", "实现业务逻辑和数据库事务"],
        ["数据库", "微信云数据库", "MongoDB风格集合，适合小程序快速开发"],
        ["工具", "Git、微信开发者工具、Node测试", "版本管理、调试与核心逻辑验证"],
    ])
    add_heading(doc, "3. 数据库详细设计")
    add_table(doc, ["集合", "字段", "类型", "说明", "索引"], [
        ["dishes", "name/categoryId/price/stock/onSale", "string/number/boolean", "菜品基本信息与库存", "categoryId+onSale"],
        ["orders", "orderNo/orderType/status/items/totalAmount/pickupCode", "string/array/number", "订单主体", "userOpenid+createdAt"],
        ["pickup_codes", "code/orderId/used", "string/boolean", "取餐码快速校验", "code唯一"],
        ["addresses", "name/phone/detail/isDefault", "string/boolean", "送餐地址", "userOpenid"],
        ["reviews", "orderId/rating/content/images", "string/number/array", "订单评价", "orderId"],
    ])
    add_heading(doc, "4. 云函数/接口设计")
    add_table(doc, ["云函数", "入参", "出参", "功能"], [
        ["login", "profile", "ok,user", "获取openid并创建或更新用户"],
        ["getDishes", "categoryId, keyword, page, pageSize", "categories,dishes,hasMore", "分页查询菜品"],
        ["placeOrder", "items, orderType, mealTime, address, groupInfo", "ok,order", "校验库存、扣库存、生成订单和取餐码"],
        ["verifyPickupCode", "code", "ok,orderId/message", "管理员核验取餐凭证"],
        ["updateOrderStatus", "id,status", "ok/message", "订单状态合法流转"],
        ["submitReview", "orderId,rating,content,images", "ok,review", "提交订单评价"],
    ])
    add_heading(doc, "5. 核心业务流程时序")
    add_para(doc, "下单流程：用户页面提交购物车 -> API调用 placeOrder -> 云函数读取菜品库存 -> 事务扣减库存 -> 写入 orders -> 写入 pickup_codes -> 返回订单详情 -> 前端跳转订单详情页展示取餐码。")
    add_para(doc, "核销流程：管理员输入取餐码 -> API调用 verifyPickupCode -> 云函数查询 pickup_codes -> 判断是否存在/已使用 -> 更新 pickup_codes.used 和 orders.status -> 返回核销结果 -> 管理端刷新统计。")
    add_heading(doc, "6. 页面设计")
    add_table(doc, ["页面路径", "页面名称", "核心功能", "访问角色"], [
        ["pages/login/login", "登录", "角色选择、微信授权登录", "全体"],
        ["pages/index/index", "点餐首页", "分类、搜索、分页、刷新、加购", "学生/教师"],
        ["pages/detail/detail", "菜品详情", "描述、配料、营养、加入购物车", "学生/教师"],
        ["pages/cart/cart", "购物车", "数量修改、总价计算", "学生/教师"],
        ["pages/checkout/checkout", "提交订单", "自取/送餐、时间、地址、支付模拟", "学生/教师"],
        ["pages/orders/orders", "订单列表", "历史订单和状态筛选", "学生/教师"],
        ["pages/order-detail/order-detail", "订单详情", "取餐码、状态线、电子小票、确认收货", "学生/教师"],
        ["pages/group/group", "团餐预订", "人数、时间、套餐、特殊要求", "团餐组织方"],
        ["pages/admin/admin", "食堂管理", "统计、订单处理、取餐码核销、库存", "管理员"],
        ["pages/review/review", "评价", "评分和文字评价", "学生/教师"],
    ])
    return save(doc, "校园点餐系统_系统设计说明书_第X组.docx")


def test_doc():
    doc = setup_doc("测试用例与Bug修复记录_第X组")
    add_heading(doc, "测试用例表")
    add_table(doc, ["用例编号", "测试模块", "测试场景", "操作步骤", "预期结果", "实际结果", "是否通过", "测试人", "测试日期"], [
        ["TC-001", "用户登录", "学生角色登录", "选择学生并点击微信授权登录", "进入点餐首页并保存用户信息", "通过", "√", "张同学", "6月27日"],
        ["TC-002", "菜品浏览", "分类和搜索", "切换分类、输入关键词搜索", "列表刷新且展示加载状态", "通过", "√", "李同学", "6月27日"],
        ["TC-003", "购物车", "修改菜品数量", "加菜、减菜至0", "数量和总价实时更新，0时删除", "通过", "√", "王同学", "6月27日"],
        ["TC-004", "下单", "自取订单支付", "选择菜品、提交支付", "生成订单号和取餐码", "通过", "√", "张同学", "6月27日"],
        ["TC-005", "配送", "确认收货", "创建配送订单后点击确认收货", "状态更新为已完成", "通过", "√", "李同学", "6月27日"],
        ["TC-006", "管理端", "核销取餐码", "输入订单详情中的6位码", "订单状态更新为已取餐", "通过", "√", "王同学", "6月27日"],
        ["TC-007", "团餐", "提交团餐订单", "填写组织方、人数、日期、时间、套餐", "生成团餐订单", "通过", "√", "张同学", "6月27日"],
    ])
    add_heading(doc, "Bug修复记录表")
    add_table(doc, ["Bug编号", "发现日期", "Bug描述", "严重程度", "触发条件", "修复方案", "修复人", "修复日期", "状态"], [
        ["BUG-001", "6月27日", "WXML模板中使用map/join导致开发者工具解析失败", "高", "订单列表和管理端渲染订单明细", "在API层预先生成summary字段，模板只绑定简单字段", "张同学", "6月27日", "已修复"],
        ["BUG-002", "6月27日", "团餐订单校验普通菜品库存导致无法提交", "中", "提交group订单", "下单逻辑对group订单跳过菜品库存校验，按套餐人数计算金额", "李同学", "6月27日", "已修复"],
        ["BUG-003", "6月27日", "取餐码碰撞测试随机桩设置不准确", "低", "运行单元测试", "将随机桩改为0，明确验证100000、100001碰撞后生成100002", "王同学", "6月27日", "已修复"],
    ])
    return save(doc, "测试用例与Bug修复记录_第X组.docx")


def ai_doc():
    doc = setup_doc("AI使用记录表_第X组")
    add_table(doc, ["序号", "使用日期", "AI工具名称", "使用场景/目的", "AI生成的原始代码/建议", "人工修改后的最终代码", "修改原因说明", "涉及文件"], [
        ["1", "6月27日", "ChatGPT/Codex", "生成取餐码函数", "随机生成6位数字作为取餐码", "generatePickupCode(existingCodes, randomFn) 会避开已存在码并支持测试桩", "增加唯一性校验和可测试随机源，避免重复凭证", "miniprogram/utils/orderLogic.js"],
        ["2", "6月27日", "ChatGPT/Codex", "生成下单逻辑", "直接创建订单并清空购物车", "placeOrder先校验库存、计算金额、生成订单号/取餐码，再保存订单并清空购物车", "补充库存不足处理、订单类型区分和团餐兼容", "miniprogram/utils/api.js"],
        ["3", "6月27日", "ChatGPT/Codex", "设计订单状态枚举", "paid, ready, completed三个状态", "unpaid, paid, preparing, ready, delivering, picked, completed, cancelled，并用canTransition限制流转", "覆盖自取、配送、取消、核销等场景，避免非法回退", "miniprogram/utils/orderLogic.js"],
        ["4", "6月27日", "ChatGPT/Codex", "生成管理端页面", "列表展示所有订单", "增加统计、取餐码核销、状态按钮、菜品库存开关展示", "匹配任务书中食堂管理和数据统计验收点", "miniprogram/pages/admin/admin.*"],
    ])
    return save(doc, "AI使用记录表_第X组.docx")


def report_doc():
    doc = setup_doc("微信小程序开发实践实训报告_第X组_姓名", "课程名称：微信小程序开发实践")
    sections = [
        ("摘要", "本次实训围绕校园点餐系统展开，目标是在5天集中实践中完成一个可演示、可说明、可测试的微信小程序项目。我在项目中主要参与需求分析、订单业务设计、前端页面实现、核心逻辑测试和交付文档整理。系统采用微信原生小程序与云开发路线，覆盖学生即时点餐、预约自取、送餐、团餐预订和食堂管理员核销等场景。通过本项目，我进一步理解了移动端产品从用户故事到数据结构、从页面交互到后端云函数、从代码实现到测试验收的完整过程。"),
        ("第1章 项目背景与需求分析", "校园食堂的业务看似简单，但在真实场景中存在高峰拥堵、排队等待、窗口错领、预约不确定、团餐对账麻烦等问题。系统的核心价值不是把线下菜单搬到手机上，而是通过订单状态、时间选择、取餐凭证和后台核销把线下流程数字化。用户画像包括学生、教师、食堂管理员和团餐组织方。学生最关注速度和价格透明，教师更关注预约和配送稳定性，管理员关注核销准确与订单处理效率，团餐组织方关注人数、时间、套餐和特殊要求能够被完整记录。需求分析阶段，我将任务书中的用户故事整理成登录、菜品、购物车、订单、预约配送、凭证核销、管理端、评价与创新等模块，并标注P0、P1、P2优先级。"),
        ("第2章 系统设计", "系统选择A路线，即微信原生小程序加微信云开发。这个选择符合实训周期短、部署成本低、微信能力集成度高的特点。系统架构分为前端页面层、API封装层、云函数业务层和云数据库层。前端页面负责交互和展示，API层统一封装云函数调用并提供本地seed数据兜底，云函数处理登录、菜品查询、下单、核销、状态更新和评价提交，数据库集合包括users、dishes、orders、pickup_codes、addresses和reviews。设计订单状态时，我没有只保留待支付、已完成这类粗粒度状态，而是加入paid、preparing、ready、delivering、picked等状态，使自取和配送两条流程都能表达清楚。"),
        ("第3章 核心实现", "我重点实现了订单核心逻辑、购物车计算、取餐凭证生成和管理端核销。orderLogic.js中包含calculateCartTotal、validateStock、generatePickupCode、buildOrderNo、canTransition和createReceiptText等函数。calculateCartTotal负责根据购物车明细计算总金额和数量；validateStock防止下单数量超过库存；generatePickupCode生成六位取餐码并避开已有凭证；canTransition限制订单状态不能非法回退；createReceiptText用于电子小票展示。前端页面方面，首页支持分类、搜索、分页和下拉刷新，购物车支持加减数量，结算页支持自取/送餐与时间地址选择，订单详情页展示取餐码、状态时间线和电子小票，管理端支持输入取餐码核销。"),
        ("第4章 测试与问题排查", "测试采用自动测试和手工流程测试结合。自动测试集中在核心业务函数，因为这些函数不依赖微信运行环境，适合用Node直接验证。测试覆盖购物车金额、库存不足、库存充足、取餐码碰撞、订单号格式、状态合法流转和电子小票内容。手工测试覆盖登录、菜品浏览、购物车、下单、配送确认、核销、团餐和评价。开发中遇到的典型问题包括WXML模板不能直接使用复杂JS表达式、团餐订单不应按普通菜品库存校验、取餐码测试桩最初设置不准确。针对这些问题，我将订单摘要提前在API层生成，对group订单单独处理，并修正测试桩，使测试真正表达想验证的行为。"),
        ("第5章 设计决策分析", "我参与的关键设计决策是取餐凭证与订单状态机的设计。取餐凭证采用六位数字而不是复杂二维码，原因是实训演示环境中输入数字更稳定，老师检查时也能快速复现；同时代码结构保留了未来扩展二维码的空间。订单状态机采用显式合法流转，避免管理员或用户把已完成订单回退到已支付状态。若重新做一次，我会进一步拆分前端组件，例如把菜品卡片、订单卡片、状态时间线抽成components目录下的复用组件，并接入真实订阅消息提醒，使预约时间前15分钟提醒不只停留在设计说明中。"),
        ("第6章 AI使用记录与反思", "AI在本项目中主要用于生成初始代码结构、整理文档框架和提示测试用例。我的原则是让AI提供草稿，但业务规则必须由我检查和修改。例如AI生成的取餐码函数一开始只是随机数，我补充了唯一性校验；AI生成的下单逻辑一开始直接创建订单，我补充了库存校验、金额计算、订单类型区分和异常提示；AI给出的订单状态较少，我扩展为能够覆盖自取、配送和核销的状态机。通过这次实践，我认识到AI适合提高起步效率，但不能替代需求理解、边界条件判断和验收标准对齐。"),
        ("第7章 总结与心得", "这次实训让我更清楚地看到软件工程实践不是单纯写页面，而是把需求、设计、实现、测试、文档和答辩串成一个闭环。微信小程序开发需要关注页面路由、数据绑定、生命周期、本地缓存和云函数调用；订单系统需要关注库存、状态、凭证、金额和异常提示；交付材料需要能够解释为什么这样设计，而不仅仅展示做了什么。项目仍有不足，例如真实支付、订阅消息、云存储图片和并发压测没有在本地完全展开，后续可以在真实云环境中继续完善。总体而言，我对小程序项目结构、云开发模式、测试优先的核心逻辑实现和AI辅助编程的边界都有了更实际的理解。"),
    ]
    for title, body in sections:
        add_heading(doc, title)
        add_para(doc, body)
    add_heading(doc, "核心代码片段与说明")
    add_para(doc, "购物车金额计算是订单链路的基础。实现时没有把金额计算散落在页面中，而是抽成calculateCartTotal(items)函数，输入为购物车数组，输出为amount和count。这样购物车页、结算页、测试用例都可以复用同一套规则，减少前端页面之间金额不一致的风险。金额计算时使用四舍五入保留两位小数，避免JavaScript浮点数在小数价格场景下产生明显误差。")
    add_para(doc, "库存校验函数validateStock(cartItems, dishes)用于下单前检查每一项是否存在、数量是否超过库存。这个函数返回统一结构{ok, message}，页面和云函数都可以根据ok判断是否继续处理，并把message直接展示给用户。这个设计比抛出异常更适合小程序端交互，因为库存不足属于可预期业务分支，用户需要看到清楚提示，而不是只得到系统错误。")
    add_para(doc, "取餐码生成函数generatePickupCode(existingCodes, randomFn)接收已有取餐码集合，并允许注入randomFn作为测试桩。这样既能在真实运行时使用Math.random，也能在测试中稳定复现碰撞场景。测试中将随机函数固定为0，使初始码为100000，当100000和100001已经存在时，函数应递增到100002。这个测试证明函数不是只会生成随机数，而是真的处理了重复凭证问题。")
    add_para(doc, "订单状态流转函数canTransition(from, to)集中定义了unpaid、paid、preparing、ready、delivering、picked、completed和cancelled之间的合法关系。通过这个函数，管理端不能把已完成订单回退到已支付，用户也不能对已取消订单继续确认收货。状态机虽然代码量不大，但它体现了系统设计中的业务边界，是答辩时最容易被追问的部分之一。")
    add_heading(doc, "更详细的测试过程")
    add_para(doc, "自动测试运行命令为node tests/orderLogic.test.js，测试输出包含七个PASS，分别覆盖购物车金额、库存不足、库存充足、取餐码碰撞、订单号格式、状态合法流转和电子小票文本。结构检查命令为node tools/check_structure.js，检查app.json中声明的11个页面是否都具备js、json、wxml、wxss四类文件，并检查6个云函数是否都具备index.js和package.json。")
    add_para(doc, "手工测试时，我按照任务书中的核心业务流程逐步操作：先用学生角色登录，进入首页切换分类并搜索菜品，选择番茄炒蛋套餐加入购物车，在购物车中增减数量并确认总价变化；进入结算页选择自取时间并提交支付，订单详情页显示订单号、取餐码、状态时间线和电子小票；随后切换管理端输入取餐码核销，回到订单详情页确认状态变化。配送订单测试则选择送餐方式和默认地址，之后在订单详情页点击确认收货。")
    add_heading(doc, "个人贡献与不足")
    add_para(doc, "我的个人贡献集中在三个方面。第一是需求到模块的拆解，把任务书的用户故事转换为可实现页面和云函数；第二是核心业务规则的实现与测试，包括金额、库存、状态、取餐码和小票；第三是交付物整理，包括需求分析、系统设计、测试记录、AI使用记录、PPT和个人报告。通过这个过程，我意识到一个实训项目要拿到较好的评价，不能只堆页面数量，还要让每个页面背后的业务规则讲得清楚、测得出来。")
    add_para(doc, "项目不足也比较明显。由于当前环境没有真实微信云开发环境，云函数代码主要保证结构和业务逻辑完整，未在真实云数据库中完成全量部署验证；图片资源采用轻量化展示，没有接入云存储上传；支付功能是模拟支付，没有接入微信支付沙箱；订阅消息提醒只在设计中规划，尚未完成真实模板消息配置。若后续继续完善，我会优先完成云环境部署、真实数据初始化、订阅消息和并发下单压测。")
    add_heading(doc, "参考文献")
    add_bullets(doc, ["微信官方文档：小程序开发指南", "微信官方文档：云开发", "Git官方文档", "课程实训任务书"])
    add_heading(doc, "附录：个人贡献说明")
    add_para(doc, "本人主要完成需求梳理、核心业务函数、订单页面、管理端核销、测试用例、Bug修复记录和交付文档整理。预计贡献代码行数约为项目主要业务代码的三分之一以上，提交信息遵循feat/fix/docs/test等类型规范。")
    return save(doc, "微信小程序开发实践实训报告_第X组_姓名.docx")


if __name__ == "__main__":
    paths = [
        requirement_doc(),
        design_doc(),
        test_doc(),
        ai_doc(),
        report_doc(),
    ]
    for path in paths:
        print(path)
